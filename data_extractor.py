from file_loader import FileLoader
import os
import fitz  # PyMuPDF
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from io import BytesIO

class DataExtractor:
    def __init__(self, file_loader: FileLoader):
        self.file_loader = file_loader

    def extract_text(self, file_path: str):
        # Implement text extraction logic
        pass

    def extract_links(self, file_path: str):
        # Implement link extraction logic
        pass

    def extract_images(self, file_path: str):
        # Implement image extraction logic
        pass

    def extract_tables(self, file_path: str):
        # Implement table extraction logic
        pass

class PDFDataExtractor(DataExtractor):
    def extract_text(self, file_path: str):
        reader, file = self.file_loader.load_file(file_path)
        text = ""
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text += page.extract_text()
        file.close()
        return text

    def extract_links(self, file_path: str):
        links = []
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            for link in page.get_links():
                if "uri" in link:
                    links.append((link["uri"], page.get_text("text", clip=link["from"])))
        return links

    def extract_images(self, file_path: str):
        images = []
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                images.append(image_bytes)
        return images

    def extract_tables(self, file_path: str):
        tables = []
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text("text")
            # Implement a simple table extraction logic
            # This is a placeholder and may need to be improved for complex tables
            if "Table" in text:
                tables.append(text)
        return tables

class DOCXDataExtractor(DataExtractor):
    def extract_text(self, file_path: str):
        doc = self.file_loader.load_file(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text
        return text

    def extract_links(self, file_path: str):
        links = []
        doc = self.file_loader.load_file(file_path)
        for para in doc.paragraphs:
            for run in para.runs:
                if run.hyperlink:
                    links.append((run.text, run.hyperlink.target))
        return links

    def extract_images(self, file_path: str):
        images = []
        doc = self.file_loader.load_file(file_path)
        for rel in doc.part.rels.values():
            if rel.reltype == RT.IMAGE:
                image_part = doc.part.related_parts[rel.target_ref]
                image_bytes = image_part.blob
                images.append(image_bytes)
        return images

    def extract_tables(self, file_path: str):
        tables = []
        doc = self.file_loader.load_file(file_path)
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables

class PPTDataExtractor(DataExtractor):
    def extract_text(self, file_path: str):
        presentation = self.file_loader.load_file(file_path)
        text = ""
        for slide in presentation.slides:
            for shape in slide.shapes:
                if hasattr(shape, "text"):
                    text += shape.text
        return text

    def extract_links(self, file_path: str):
        # Implement PPT link extraction logic
        pass

    def extract_images(self, file_path: str):
        images = []
        presentation = self.file_loader.load_file(file_path)
        for slide in presentation.slides:
            for shape in slide.shapes:
                if shape.shape_type == 13:  # Picture shape type
                    image = shape.image
                    images.append(image)
        return images

    def extract_tables(self, file_path: str):
        tables = []
        presentation = self.file_loader.load_file(file_path)
        for slide in presentation.slides:
            for shape in slide.shapes:
                if shape.has_table:
                    table = shape.table
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append(table_data)
        return tables
